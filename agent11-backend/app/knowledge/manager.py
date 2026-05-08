"""知识库管理 - 文件系统 + ChromaDB 向量搜索"""
from __future__ import annotations

import os
import shutil
from pathlib import Path
from typing import Optional
import structlog

from app.config import get_settings
from app.knowledge.chromadb import ChromaDBClient

logger = structlog.get_logger()


class KnowledgeManager:
    """知识库管理器 - 管理文件系统上的 .md 文件和 ChromaDB 向量索引"""

    def __init__(self):
        settings = get_settings()
        self.base_path = Path(settings.knowledge_base_path)
        self.base_path.mkdir(parents=True, exist_ok=True)

    def _get_file_path(self, path: str) -> Path:
        """获取文件的完整路径"""
        clean_path = path.lstrip('/')
        return self.base_path / clean_path

    async def list_files(self, path: str = "") -> list[dict]:
        """列出指定目录下的文件和文件夹"""
        target_path = self._get_file_path(path)

        if target_path.exists() and target_path.is_file():
            return [self._file_to_dict(target_path)]

        items = []
        try:
            for item in sorted(target_path.iterdir()):
                items.append(self._item_to_dict(item))
        except FileNotFoundError:
            return []

        return items

    def _item_to_dict(self, item: Path, parent_path: str = "") -> dict:
        """将路径项转为字典"""
        if parent_path:
            rel_path = Path(parent_path) / item.name
        else:
            rel_path = item.relative_to(self.base_path)
        stat = item.stat()

        if item.is_dir():
            return {
                "filename": item.name,
                "path": str(rel_path),
                "isFolder": True,
                "createdAt": None,
                "updatedAt": None,
            }
        else:
            return {
                "filename": item.name,
                "path": str(rel_path),
                "isFolder": False,
                "createdAt": stat.st_ctime,
                "updatedAt": stat.st_mtime,
            }

    def _file_to_dict(self, path: Path) -> dict:
        """将文件转为字典"""
        rel_path = path.relative_to(self.base_path)
        stat = path.stat()

        return {
            "filename": path.name,
            "path": str(rel_path),
            "isFolder": False,
            "createdAt": stat.st_ctime,
            "updatedAt": stat.st_mtime,
        }

    async def get_file_content(self, filename: str) -> Optional[str]:
        """读取文件内容"""
        file_path = self._get_file_path(filename)

        if not file_path.exists() or file_path.is_dir():
            return None

        try:
            return file_path.read_text(encoding="utf-8")
        except Exception as e:
            logger.error("read_file_failed", path=str(file_path), error=str(e))
            return None

    async def create_folder(self, name: str, parent_path: str = "") -> bool:
        """创建文件夹"""
        folder_path = self._get_file_path(parent_path) / name

        try:
            folder_path.mkdir(parents=True, exist_ok=True)
            return True
        except Exception as e:
            logger.error("create_folder_failed", path=str(folder_path), error=str(e))
            return False

    async def delete_folder(self, path: str) -> bool:
        """删除文件夹"""
        folder_path = self._get_file_path(path)

        if not folder_path.exists() or not folder_path.is_dir():
            return False

        try:
            shutil.rmtree(folder_path)
            return True
        except Exception as e:
            logger.error("delete_folder_failed", path=str(folder_path), error=str(e))
            return False

    async def delete_file(self, filename: str) -> bool:
        """删除文件"""
        file_path = self._get_file_path(filename)

        if not file_path.exists() or file_path.is_dir():
            return False

        try:
            file_path.unlink()
            return True
        except Exception as e:
            logger.error("delete_file_failed", path=str(file_path), error=str(e))
            return False

    async def save_file(self, filename: str, content: str) -> bool:
        """保存文件内容"""
        file_path = self._get_file_path(filename)

        try:
            file_path.parent.mkdir(parents=True, exist_ok=True)
            file_path.write_text(content, encoding="utf-8")

            await self._index_file_to_chroma(file_path)
            return True
        except Exception as e:
            logger.error("save_file_failed", path=str(file_path), error=str(e))
            return False

    async def upload_file(self, filename: str, content: bytes, folder_path: str = "") -> tuple[bool, str]:
        """上传文件，自动转换格式"""
        file_path = self._get_file_path(folder_path) / filename

        try:
            file_path.parent.mkdir(parents=True, exist_ok=True)

            ext = filename.lower().split('.')[-1]

            if ext in ('xls', 'xlsx'):
                text_content = await self._convert_excel_to_md(content, filename)
            elif ext in ('doc', 'docx'):
                text_content = await self._convert_word_to_md(content, filename)
            elif ext == 'pdf':
                text_content = await self._convert_pdf_to_text(content, filename)
            else:
                text_content = content.decode('utf-8', errors='ignore')

            md_filename = filename.rsplit('.', 1)[0] + '.md'
            md_path = file_path.parent / md_filename

            md_path.write_text(text_content, encoding='utf-8')

            await self._index_file_to_chroma(md_path)

            return True, md_filename
        except Exception as e:
            logger.error("upload_file_failed", filename=filename, error=str(e))
            return False, str(e)

    async def _convert_excel_to_md(self, content: bytes, filename: str) -> str:
        """将 Excel 转换为 Markdown 表格"""
        import io
        try:
            import pandas as pd
        except ImportError:
            return f"# {filename}\n\nExcel file content (pandas not available)"
        try:
            df = pd.read_excel(io.BytesIO(content))
            lines = ["| " + " | ".join(str(col) for col in df.columns) + " |"]
            lines.append("|" + "|".join(["---"] * len(df.columns)) + "|")
            for _, row in df.iterrows():
                lines.append("| " + " | ".join(str(val) for val in row) + " |")
            return "\n".join(lines)
        except Exception as e:
            logger.warning("excel_convert_failed", filename=filename, error=str(e))
            return f"# {filename}\n\nFailed to convert Excel: {str(e)}"

    async def _convert_word_to_md(self, content: bytes, filename: str) -> str:
        """将 Word 转换为 Markdown"""
        import io
        try:
            from docx import Document
        except ImportError:
            return f"# {filename}\n\nWord document content (preview)\n\nNote: Install python-docx for full conversion"
        try:
            doc = Document(io.BytesIO(content))
            lines = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # 处理标题
                    if para.style.name.startswith('Heading'):
                        level = para.style.name.split(' ')[-1]
                        try:
                            level = int(level)
                            lines.append(f"{'#' * (level + 1)} {text}")
                        except:
                            lines.append(f"## {text}")
                    else:
                        lines.append(text)
            # 处理表格
            for table in doc.tables:
                rows_data = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows_data.append(cells)
                if rows_data:
                    lines.append("")
                    lines.append("| " + " | ".join(rows_data[0]) + " |")
                    lines.append("|" + "|".join(["---"] * len(rows_data[0])) + "|")
                    for row in rows_data[1:]:
                        lines.append("| " + " | ".join(row) + " |")
            return "\n".join(lines) if lines else f"# {filename}\n\n(Empty document)"
        except Exception as e:
            logger.warning("word_convert_failed", filename=filename, error=str(e))
            return f"# {filename}\n\nFailed to convert Word: {str(e)}"

    async def _convert_pdf_to_text(self, content: bytes, filename: str) -> str:
        """将 PDF 转换为文本"""
        return f"# {filename}\n\nPDF document content (preview)\n\nNote: Install PyPDF2 for full text extraction"

    async def _index_file_to_chroma(self, file_path: Path) -> None:
        """将文件索引到 ChromaDB"""
        try:
            content = file_path.read_text(encoding="utf-8")
            chroma = ChromaDBClient.get_instance()
            if chroma:
                await chroma.add_to_collection(
                    collection_name="user_knowledge",
                    documents=[content],
                    metadata=[{"filename": file_path.name, "path": str(file_path)}]
                )
        except Exception as e:
            logger.error("index_file_failed", path=str(file_path), error=str(e))

    async def search(self, query: str, limit: int = 5) -> list[dict]:
        """语义搜索知识库"""
        chroma = ChromaDBClient.get_instance()
        if not chroma:
            return []

        try:
            results = await chroma.query(
                collection_name="user_knowledge",
                query=query,
                n_results=limit
            )

            search_results = []
            if results.get("ids") and results["ids"]:
                for i, doc_id in enumerate(results["ids"][0]):
                    search_results.append({
                        "filename": results.get("metadatas", [[{}]])[0][i].get("filename", "unknown"),
                        "content": results.get("documents", [[]])[0][i] if results.get("documents") else "",
                        "score": results.get("distances", [[]])[0][i] if results.get("distances") else 1.0,
                    })

            return search_results
        except Exception as e:
            logger.error("search_failed", query=query, error=str(e))
            return []

    async def reindex_all(self) -> int:
        """重建所有文件的索引"""
        count = 0
        try:
            for md_file in self.base_path.rglob("*.md"):
                await self._index_file_to_chroma(md_file)
                count += 1
        except Exception as e:
            logger.error("reindex_failed", error=str(e))
        return count